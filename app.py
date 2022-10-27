"""
Simple app to upload an image via a web form 
and view the inference results on the image in the browser.
"""
import argparse
import io
import math
import os
from PIL import Image
from flask_cors import CORS
import numpy as np
 
from base64 import b64encode

import base64
import torch
import logging
#import azure.functions as func
import tempfile
from os import listdir

import pandas as pd

import openpyxl
from flask import Flask, render_template, request, redirect,jsonify,send_file
from subprocess import STDOUT, check_call , call
#from w3lib.url import parse_data_uri

import os
import json

from docx import Document # for pdf format
from docx.shared import Pt # for pdf format
from docx.shared import Inches
#from docx2pdf import convert # for pdf format
#from msoffice2pdf import convert #for pdf format
#import pythoncom # for pdf format
app = Flask(__name__)
CORS(app)
check_call(['apt-get', 'update'], stdout=open(os.devnull,'wb'), stderr=STDOUT)
check_call(['apt-get', 'install', '-y', 'libgl1'], stdout=open(os.devnull,'wb'), stderr=STDOUT)
check_call(['apt-get', 'install', '-y', 'libglib2.0-0'], stdout=open(os.devnull,'wb'), stderr=STDOUT)
check_call([ 'apt-get', 'update','-y'], stdout=open(os.devnull,'wb'), stderr=STDOUT)
check_call([ 'apt-get', 'install' ,'-y','abiword'], stdout=open(os.devnull,'wb'), stderr=STDOUT)


github='ultralytics/yolov5'
torch.hub.list(github, trust_repo=True)
model = torch.hub.load("ultralytics/yolov5", "custom", path = "./rings18.pt", force_reload=True)
  
model.classes=[3 ,10,11 ,12, 17]





@app.route("/", methods=["GET", "POST"])
def predict():
    print("Here")
    data = request.get_json(force=True)
    data_keys=list(data.keys())
    cyl_index=-1
    defect_df_all_cyl={}
    img_list=[]
    for cyl in  data.values():
        cyl_index=cyl_index+1
        print(cyl,type(cyl))

#---
   
        header, encoded = cyl. split(",", 1)
        data = base64.b64decode(encoded)


        with open("image.png", "wb") as f:
            f.write(data)   
     

        img = Image.open(io.BytesIO(data))
        results = model(img, size=640)
        img = np.squeeze(results.render())
        # datatoexcel = pd.ExcelWriter('results.xlsx')
        # results.to_excel(datatoexcel)
        # datatoexcel.save()

        
        print("RESULT=======",results)
        file_object = io.BytesIO()
        
        data = Image.fromarray(img)
        data.save(file_object, 'JPEG')
        
        base64img = "data:image/png;base64,"+b64encode(file_object.getvalue()).decode('ascii')



 
        res_tensor=results.xyxy[0]  # im1 predictions (tensor)
        h=data.height

        img_list.append(file_object) ########### adding images in docx #########################
        print(results.pandas().xyxy[0] ) # im1 predictions (pandas)
        print("y ",res_tensor[0][1])
        print("c ",int(res_tensor[0][5]))
        print("tensor len",len(res_tensor))
        
        rings=[]
        for i in range(0,len(res_tensor)):
            
            print("percent=====",res_tensor[i][1]/h)
            if res_tensor[i][1]/h <=.25 :
                rings.append({"1":int(res_tensor[i][5])})
            elif res_tensor[i][1]/h <=.45 :
                rings.append({"2":int(res_tensor[i][5])})
            elif res_tensor[i][1]/h <=.75 :
                rings.append({"3":int(res_tensor[i][5])})
            elif res_tensor[i][1]/h >.75 :
                rings.append({"4":int(res_tensor[i][5])})

                     
                
        
        def_section_brk=set()
        def_section_lub1=set()
        def_section_surf=set()
        def_section_dep=set()
        def_section_lub2=set()
        def_section_brk_ls={}
        
        def_section_lub_ls={}
        def_section_surf_ls={}
        def_section_dep_ls={}
        try:
            for ring_no in range(1,5):
                
                def_section_lub_ls["Ring"+str(ring_no)]="*"
        except Exception :
            print("Excepetion")

        try:
            for ring_no in range(1,5):
                
                
                def_section_surf_ls["Ring"+str(ring_no)]="*"
        except Exception :
            print("Excepetion")
        try:
            for ring_no in range(1,5):
                
                
                def_section_dep_ls["Ring"+str(ring_no)]="*"
                
        except Exception :
            print("Excepetion")
        try:
            for ring_no in range(1,5):
               
                def_section_brk_ls["Ring"+str(ring_no)]="*"
               
        except Exception :
            print("Excepetion")

        for ring in rings:
            print(ring.values())
        
            if(list(ring.values())[0]==3):# if collapsed
                def_section_brk.add(list(ring.keys())[0]) # assign ring number
            if(list(ring.values())[0]==12): 
                def_section_surf.add(list(ring.keys())[0])
            if(list(ring.values())[0]==11 ):
                def_section_lub1.add(list(ring.keys())[0])
            if(list(ring.values())[0]==17):
                def_section_dep.add(list(ring.keys())[0])
            if(list(ring.values())[0]==10):
                def_section_lub2.add(list(ring.keys())[0])
       
        for brk in def_section_brk :
            print({"Ring"+brk:"C"})
            def_section_brk_ls.update({"Ring"+brk:"C"})
        for brk in def_section_surf :

            def_section_surf_ls.update({"Ring"+brk:"S"})
        for brk in def_section_dep :
            print({"Ring"+brk:"LC"})
            def_section_dep_ls.update({"Ring"+brk:"LC"})

        for brk in def_section_lub1 :
            def_section_lub_ls.update({"Ring"+brk:"OB"})
        for brk in def_section_lub2 :
            print({"Ring"+brk:"OB"})
            if brk not in(list(def_section_lub1)):
                def_section_lub_ls.update({"Ring"+brk:"O"})
        print(def_section_dep_ls)
       

        print("def_section_lub_ls", def_section_lub_ls)
        #js_data=results.pandas().xyxy[0].to_json(orient="records")
        
        defect_df=  {"lubrication":def_section_lub_ls, "surface":def_section_surf_ls,"deposits":def_section_dep_ls,"breakage":def_section_brk_ls
        #, "image":base64img
        }
        defect_df_all_cyl["cylinder"+(data_keys[cyl_index])]=defect_df
        # #print(defect_df)
        # print("------------------defect-------------------------------")
        
    print(defect_df_all_cyl)

    selection_lubrication=[]
    selection_surface=[]
    selection_deposits=[]
    selection_brekage=[]
    
    user_data=[]
    
    cyls=defect_df_all_cyl.keys()
    print(cyls)
    for cyl in cyls:
        print(type(cyl))
        if cyl.startswith("cylinder"):

            selection_lubrication.append(list(defect_df_all_cyl[cyl]['lubrication'].values()))
    

            selection_surface.append(list(defect_df_all_cyl[cyl]['surface'].values()))
    

            selection_brekage.append(list(defect_df_all_cyl[cyl]['breakage'].values()))
    

            selection_deposits.append(list(defect_df_all_cyl[cyl]['deposits'].values()))
 
   
    pred_per_cyl_lubrication_rev = [[selection_lubrication[j][i] for j in range(len(selection_lubrication))] for i in range(len(selection_lubrication[0]))]
    pred_per_cyl_surface_rev = [[selection_surface[j][i] for j in range(len(selection_surface))] for i in range(len(selection_surface[0]))]
    pred_per_cyl_deposits_rev = [[selection_deposits[j][i] for j in range(len(selection_deposits))] for i in range(len(selection_deposits[0]))]
    pred_per_cyl_breakage_rev = [[selection_brekage[j][i] for j in range(len(selection_brekage))] for i in range(len(selection_brekage[0]))]
    print(pred_per_cyl_lubrication_rev)



############################################################### PDF ##################################################

    print(len(data_keys))
    doc = Document('sample_report.docx')
#     doc.tables #a list of all tables in document
   
#     doc.tables[0].cell(0, 1).text = user_data["vesselName"] # vessel name
#     doc.tables[0].cell(0, 1).paragraphs[0].runs[0].font.size = Pt(7)
#     doc.tables[0].cell(0, 1).paragraphs[0].runs[0].font.name = 'Arial'
#     doc.tables[0].cell(0, 1).paragraphs[0].runs[0].font.bold = True

#     doc.tables[0].cell(0, 3).text = user_data["hullNumber"] # hull
#     doc.tables[0].cell(0, 3).paragraphs[0].runs[0].font.size = Pt(7)
#     doc.tables[0].cell(0, 3).paragraphs[0].runs[0].font.name = 'Arial'
#     doc.tables[0].cell(0, 3).paragraphs[0].runs[0].font.bold = True

#     doc.tables[0].cell(0, 5).text=user_data["vesselType"] # vessel type
#     doc.tables[0].cell(0,5).paragraphs[0].runs[0].font.size = Pt(7)
#     doc.tables[0].cell(0,5).paragraphs[0].runs[0].font.name = 'Arial'
#     doc.tables[0].cell(0,5).paragraphs[0].runs[0].font.bold = True

#     doc.tables[2].cell(0, 1).text = user_data["manufacture"] # make
#     doc.tables[2].cell(0, 1).paragraphs[0].runs[0].font.size = Pt(7)
#     doc.tables[2].cell(0, 1).paragraphs[0].runs[0].font.name = 'Arial'
#     doc.tables[2].cell(0, 1).paragraphs[0].runs[0].font.bold = True
    
#     doc.tables[4].cell(1, 1).text = str(user_data["totalRunningHours"]) #  Total running hours
#     doc.tables[4].cell(1, 1).paragraphs[0].runs[0].font.size = Pt(7)
#     doc.tables[4].cell(1, 1).paragraphs[0].runs[0].font.name = 'Arial'


#     doc.tables[4].cell(2, 1).text = user_data["cyloiltype"] # Cyl oil type
#     doc.tables[4].cell(2, 1).paragraphs[0].runs[0].font.size = Pt(7)
#     doc.tables[4].cell(2, 1).paragraphs[0].runs[0].font.name = 'Arial'
   

#     doc.tables[4].cell(3, -1).text = user_data["cyloilconsump"] # Cyl oil rate
#     doc.tables[4].cell(3, -1).paragraphs[0].runs[0].font.size = Pt(7)
#     doc.tables[4].cell(3, -1).paragraphs[0].runs[0].font.name = 'Arial'
   

    
#     doc.tables[4].cell(3, 1).text = str(user_data["normalserviceloadpercentMCR"]) # Normal service MCR
#     doc.tables[4].cell(3, 1).paragraphs[0].runs[0].font.size = Pt(7)
#     doc.tables[4].cell(3, 1).paragraphs[0].runs[0].font.name = 'Arial'

    
   
    list_ind=[2,3,-2,-1]
    doc.tables[6].style ='TableGrid' # Deposits section --tables[6]
    print("selection_deposits=",selection_deposits)
    for i in range(len(selection_deposits)):
        row_cells = doc.tables[6].add_row().cells
        ind=0
        row_cells[0].text = str(i+1)
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[0].paragraphs[0].runs[0].font.size = Pt(7)
        row_cells[0].paragraphs[0].runs[0].font.name = 'Arial'
        for j in list_ind:
            row_cells[j].text = 'l' if selection_deposits[i][ind]=='*' else selection_deposits[i][ind]
           # row_cells[0].paragraphs[0].runs[0].font.bold = True
            row_cells[j].paragraphs[0].runs[0].font.size = Pt(7)
            row_cells[j].paragraphs[0].runs[0].font.name = 'Arial'
            ind=ind+1
        row_cells[1].text = 'l' 
        # row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[1].paragraphs[0].runs[0].font.size = Pt(7)
        row_cells[1].paragraphs[0].runs[0].font.name = 'Arial'
    #Mean row
    row_cells = doc.tables[6].add_row().cells
    row_cells[0].text = "Mean"
    row_cells[0].paragraphs[0].runs[0].font.bold = True
    row_cells[0].paragraphs[0].runs[0].font.size = Pt(7)
    row_cells[0].paragraphs[0].runs[0].font.name = 'Arial'
    # Breakage section --tables[7]
    list_ind=[1,2,3,-2]
    doc.tables[7].style ='TableGrid' 
    for i in range(len(selection_brekage)):
        row_cells = doc.tables[7].add_row().cells
        ind=0
        row_cells[0].text = str(i+1)
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[0].paragraphs[0].runs[0].font.size = Pt(7)
        row_cells[0].paragraphs[0].runs[0].font.name = 'Arial'
        for j in list_ind:
            row_cells[j].text ='l' if selection_brekage[i][ind]=='*' else selection_brekage[i][ind]
           # row_cells[0].paragraphs[0].runs[0].font.bold = True
            row_cells[j].paragraphs[0].runs[0].font.size = Pt(7)
            row_cells[j].paragraphs[0].runs[0].font.name = 'Arial'
            ind=ind+1
        
        row_cells[-1].text = 'l' 
        # row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[-1].paragraphs[0].runs[0].font.size = Pt(7)
        row_cells[-1].paragraphs[0].runs[0].font.name = 'Arial'

    row_cells = doc.tables[7].add_row().cells
    row_cells[0].text = "Mean"
    row_cells[0].paragraphs[0].runs[0].font.bold = True
    row_cells[0].paragraphs[0].runs[0].font.size = Pt(7)
    row_cells[0].paragraphs[0].runs[0].font.name = 'Arial'

# Surface section --tables[10]
    list_ind=[1,2,3,4]
    list_ind_extra=[5,6,7,-2,-1]
    doc.tables[10].style ='TableGrid' 
    for i in range(len(selection_surface)):
        row_cells = doc.tables[10].add_row().cells
        ind=0
        row_cells[0].text = str(i+1)
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[0].paragraphs[0].runs[0].font.size = Pt(7)
        row_cells[0].paragraphs[0].runs[0].font.name = 'Arial'
        for j in list_ind:
            row_cells[j].text = 'Cl' if selection_surface[i][ind]=='*' else selection_surface[i][ind]
           # row_cells[0].paragraphs[0].runs[0].font.bold = True
            row_cells[j].paragraphs[0].runs[0].font.size = Pt(7)
            row_cells[j].paragraphs[0].runs[0].font.name = 'Arial'
            ind=ind+1
        for ext in list_ind_extra:
            row_cells[ext].text = 'Cl' 
           # row_cells[0].paragraphs[0].runs[0].font.bold = True
            row_cells[ext].paragraphs[0].runs[0].font.size = Pt(7)
            row_cells[ext].paragraphs[0].runs[0].font.name = 'Arial'
        
    row_cells = doc.tables[10].add_row().cells #Mean row
    row_cells[0].text = "Mean"
    row_cells[0].paragraphs[0].runs[0].font.bold = True
    row_cells[0].paragraphs[0].runs[0].font.size = Pt(7)
    row_cells[0].paragraphs[0].runs[0].font.name = 'Arial'

# Lubrication section --tables[11]
    list_ind=[1,2,3,4]
    list_ind_extra=[5,6,-2,-1]
    doc.tables[11].style ='TableGrid' 
    for i in range(len(selection_lubrication)):
        row_cells = doc.tables[11].add_row().cells
        ind=0
        row_cells[0].text = str(i+1)
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        row_cells[0].paragraphs[0].runs[0].font.size = Pt(7)
        row_cells[0].paragraphs[0].runs[0].font.name = 'Arial'
        for j in list_ind:
            row_cells[j].text = 'N' if selection_lubrication[i][ind]=='*' else selection_lubrication[i][ind]
           # row_cells[0].paragraphs[0].runs[0].font.bold = True
            row_cells[j].paragraphs[0].runs[0].font.size = Pt(7)
            row_cells[j].paragraphs[0].runs[0].font.name = 'Arial'
            ind=ind+1
        
        for ext in list_ind_extra: # Ring5 , Piston Skirt,Rod,Liner
            row_cells[ext].text = 'N' 
           # row_cells[0].paragraphs[0].runs[0].font.bold = True
            row_cells[ext].paragraphs[0].runs[0].font.size = Pt(7)
            row_cells[ext].paragraphs[0].runs[0].font.name = 'Arial'

    row_cells = doc.tables[11].add_row().cells  #Mean row
    row_cells[0].text = "Mean"
    row_cells[0].paragraphs[0].runs[0].font.bold = True
    row_cells[0].paragraphs[0].runs[0].font.size = Pt(7)
    row_cells[0].paragraphs[0].runs[0].font.name = 'Arial'

########### adding images #########################
    
 
    
    
    table = doc.add_table(rows=math.ceil(len(img_list)/2), cols=2)
    table.style ='TableGrid' 
    #for im in img_list:
    im_count=0
    for row in table.rows:
        
        for cell in row.cells:
            im_count+=1
            
            if im_count<= len(img_list): # odd number of cylinders
                cell.text = "Cylinder "+str(im_count)
                p=cell.add_paragraph()
                r = p.add_run()
                r.add_picture(img_list[im_count-1],width=Inches(3))
   
        
   
        
        


    doc.save("report.docx")
    
    args = ["abiword", "--to", "report_output.pdf", "report.docx"  ]

    call(args )
        

    return "please upload an image"



if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Flask app exposing yolov5 models")
    parser.add_argument("--port", default=8000, type=int, help="port number")
    args = parser.parse_args()
    
   # detect.run(weights='yolov5s.pt', save_txt= True)
    app.run(host="0.0.0.0", port=args.port)  # debug=True causes Restarting with stat
